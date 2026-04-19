"""Tool: export_report — generate a Word report from scenario analysis."""

import agent.config  # noqa: F401
from agent.config import OUTPUT_DIR, ensure_dirs
from agent.tools.registry import register_tool


def _export_report(
    magnitude: float = 6.7,
    lat: float = 34.213,
    lon: float = -118.537,
    depth_km: float = 10.0,
    fault_type: str = "reverse",
    n_bridges: int = 100,
    radius_km: float = 30.0,
    n_realizations: int = 30,
    seed: int = 42,
    title: str = "Earthquake Bridge Risk Assessment Report",
) -> str:
    """Generate a Word document report with scenario analysis results."""
    import math
    import numpy as np
    from datetime import datetime
    from pathlib import Path

    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from src.hazard import EarthquakeScenario, compute_sa_at_sites
    from src.exposure import (
        generate_synthetic_portfolio, portfolio_to_sites, portfolio_summary,
    )
    from src.engine import run_deterministic
    from src.loss import compute_bridge_loss
    from src.fragility import damage_state_probabilities
    from src.hazus_params import DAMAGE_STATE_ORDER

    ensure_dirs()

    # ── Run analysis ──
    scenario = EarthquakeScenario(
        Mw=magnitude, lat=lat, lon=lon,
        depth_km=depth_km, fault_type=fault_type,
    )
    portfolio = generate_synthetic_portfolio(
        n_bridges=n_bridges, center=(lat, lon),
        radius_km=radius_km, seed=seed,
    )
    result = run_deterministic(
        scenario=scenario, portfolio=portfolio,
        n_realizations=n_realizations, seed=seed,
    )

    if not result.loss_results:
        return "Error: no loss results generated"

    losses = np.array([r.total_loss for r in result.loss_results])
    ps = result.portfolio_summary
    total_rc = ps.get("total_replacement_cost", 1)

    # Per-bridge median Sa results
    sites = portfolio_to_sites(portfolio)
    median_sa, _ = compute_sa_at_sites(scenario, sites)

    # ── Build Word document ──
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title
    t = doc.add_heading(level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    # 1. Scenario Parameters
    doc.add_heading("1. Earthquake Scenario", level=1)
    param_table = doc.add_table(rows=6, cols=2)
    param_table.style = "Light Shading Accent 1"
    params = [
        ("Magnitude", f"M{magnitude}"),
        ("Epicenter", f"({lat:.3f}\u00b0N, {abs(lon):.3f}\u00b0W)"),
        ("Depth", f"{depth_km} km"),
        ("Fault Type", fault_type.replace("_", " ").title()),
        ("Portfolio", f"{n_bridges} bridges (synthetic, {radius_km}km radius)"),
        ("Realizations", str(n_realizations)),
    ]
    for i, (label, value) in enumerate(params):
        param_table.rows[i].cells[0].text = label
        param_table.rows[i].cells[1].text = value
        for cell in param_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # 2. Loss Summary
    doc.add_heading("2. Loss Summary", level=1)

    mean_loss = float(np.mean(losses))
    std_loss = float(np.std(losses))
    p5 = float(np.percentile(losses, 5))
    p50 = float(np.percentile(losses, 50))
    p95 = float(np.percentile(losses, 95))

    loss_table = doc.add_table(rows=7, cols=2)
    loss_table.style = "Light Shading Accent 1"
    loss_data = [
        ("Total Replacement Cost", f"${total_rc:,.0f}"),
        ("Mean Expected Loss", f"${mean_loss:,.0f}"),
        ("Median Loss (P50)", f"${p50:,.0f}"),
        ("90% Confidence Interval", f"${p5:,.0f} \u2013 ${p95:,.0f}"),
        ("Standard Deviation", f"${std_loss:,.0f}"),
        ("Mean Loss Ratio", f"{mean_loss/total_rc:.4f} ({mean_loss/total_rc*100:.2f}%)"),
        ("P95 Loss Ratio", f"{p95/total_rc:.4f} ({p95/total_rc*100:.2f}%)"),
    ]
    for i, (label, value) in enumerate(loss_data):
        loss_table.rows[i].cells[0].text = label
        loss_table.rows[i].cells[1].text = value
        for cell in loss_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # 3. Damage Distribution
    doc.add_heading("3. Damage State Distribution", level=1)
    doc.add_paragraph(
        "Expected number of bridges in each damage state (mean across "
        f"{n_realizations} realizations):"
    )

    last_r = result.loss_results[-1]
    ds_table = doc.add_table(rows=6, cols=3)
    ds_table.style = "Light Shading Accent 1"
    ds_table.rows[0].cells[0].text = "Damage State"
    ds_table.rows[0].cells[1].text = "Expected Count"
    ds_table.rows[0].cells[2].text = "Percentage"
    for cell in ds_table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    actual_n = result.portfolio_size
    for i, ds in enumerate(["none"] + DAMAGE_STATE_ORDER):
        cnt = last_r.count_by_ds.get(ds, 0)
        pct = cnt / actual_n * 100 if actual_n > 0 else 0
        ds_table.rows[i + 1].cells[0].text = ds.capitalize()
        ds_table.rows[i + 1].cells[1].text = f"{cnt:.1f}"
        ds_table.rows[i + 1].cells[2].text = f"{pct:.1f}%"
        for cell in ds_table.rows[i + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # 4. Loss by Bridge Class
    doc.add_heading("4. Loss by Bridge Class", level=1)
    class_losses = sorted(
        last_r.loss_by_class.items(), key=lambda x: x[1], reverse=True
    )
    cls_table = doc.add_table(rows=len(class_losses) + 1, cols=3)
    cls_table.style = "Light Shading Accent 1"
    cls_table.rows[0].cells[0].text = "HWB Class"
    cls_table.rows[0].cells[1].text = "Expected Loss"
    cls_table.rows[0].cells[2].text = "% of Total"
    for cell in cls_table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    for i, (cls, loss) in enumerate(class_losses):
        pct = loss / last_r.total_loss * 100 if last_r.total_loss > 0 else 0
        cls_table.rows[i + 1].cells[0].text = cls
        cls_table.rows[i + 1].cells[1].text = f"${loss:,.0f}"
        cls_table.rows[i + 1].cells[2].text = f"{pct:.1f}%"
        for cell in cls_table.rows[i + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # 5. Ground Motion
    doc.add_heading("5. Ground Motion Statistics", level=1)
    gm_table = doc.add_table(rows=3, cols=2)
    gm_table.style = "Light Shading Accent 1"
    gm_data = [
        ("Mean Median Sa(1.0s)", f"{float(np.mean(median_sa)):.4f} g"),
        ("Sa Range", f"{float(np.min(median_sa)):.4f} \u2013 {float(np.max(median_sa)):.4f} g"),
        ("GMPE Model", "Boore & Atkinson (2008)"),
    ]
    for i, (label, value) in enumerate(gm_data):
        gm_table.rows[i].cells[0].text = label
        gm_table.rows[i].cells[1].text = value

    # 6. Methodology
    doc.add_heading("6. Methodology", level=1)
    doc.add_paragraph(
        "This analysis uses the CAT411 framework implementing FEMA Hazus 6.1 "
        "earthquake bridge loss estimation methodology. Key components:"
    )
    methodology_bullets = [
        "Ground Motion: Boore & Atkinson (2008) GMPE for Sa(1.0s), with "
        "spatially correlated ground motion fields via Jayaram-Baker (2009) model.",
        "Fragility: Lognormal CDF model with parameters from Hazus Table 7.9 "
        "(28 HWB bridge classes, uniform beta = 0.6).",
        "Loss: Hazus Table 7.11 damage ratios (slight=3%, moderate=8%, "
        "extensive=25%, complete=100% of replacement cost).",
        f"Uncertainty: {n_realizations} Monte Carlo realizations of spatially "
        "correlated ground motion fields. Reported confidence intervals reflect "
        "aleatory variability only; epistemic uncertainty (GMPE model choice, "
        "fragility parameter uncertainty) is not included.",
    ]
    for b in methodology_bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Reference: FEMA (2024). Hazus 6.1 Earthquake Model Technical Manual."
    )

    # Save
    report_path = str(OUTPUT_DIR / "risk_assessment_report.docx")
    doc.save(report_path)

    return (
        f"Report saved to: {report_path}\n\n"
        f"Contents:\n"
        f"  1. Earthquake Scenario (M{magnitude}, {fault_type})\n"
        f"  2. Loss Summary (mean ${mean_loss:,.0f}, 90% CI ${p5:,.0f}-${p95:,.0f})\n"
        f"  3. Damage State Distribution\n"
        f"  4. Loss by Bridge Class\n"
        f"  5. Ground Motion Statistics\n"
        f"  6. Methodology & References"
    )


register_tool(
    name="export_report",
    description=(
        "Generate a professional Word document (.docx) report containing "
        "earthquake scenario parameters, loss summary with confidence intervals, "
        "damage state distribution, loss by bridge class, ground motion statistics, "
        "and methodology description with references. Suitable for official "
        "documents and FEMA applications."
    ),
    parameters={
        "type": "object",
        "properties": {
            "magnitude": {"type": "number", "default": 6.7},
            "lat": {"type": "number", "default": 34.213},
            "lon": {"type": "number", "default": -118.537},
            "depth_km": {"type": "number", "default": 10.0},
            "fault_type": {
                "type": "string",
                "enum": ["reverse", "strike_slip", "normal"],
                "default": "reverse",
            },
            "n_bridges": {"type": "integer", "default": 100},
            "radius_km": {"type": "number", "default": 30.0},
            "n_realizations": {"type": "integer", "default": 30},
            "seed": {"type": "integer", "default": 42},
            "title": {
                "type": "string",
                "description": "Report title",
                "default": "Earthquake Bridge Risk Assessment Report",
            },
        },
        "required": [],
    },
    function=_export_report,
)
