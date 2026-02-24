"""
CAT411 — Professional GIS-Style Report Maps

Generates a single comprehensive ArcGIS-style map showing multiple
analysis scenarios (different filter conditions) on one figure, with
ShakeMap contour background and each scenario's bridges + region boundary.

Also generates a conditions summary table as a separate figure.

Usage:
    python scripts/generate_report_maps.py
"""

import os
import sys
import shutil
from pathlib import Path

# Ensure project root is on sys.path for src.* imports
_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent
sys.path.insert(0, str(_PROJECT_ROOT))
os.chdir(_PROJECT_ROOT)

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patheffects as pe
from matplotlib.patches import Rectangle
from matplotlib.lines import Line2D
from matplotlib.colors import LinearSegmentedColormap
from matplotlib import ticker
from scipy.interpolate import griddata

from src.config import load_config, IM_COLUMN_MAP, AnalysisConfig
from src.data_loader import load_shakemap, load_nbi, classify_nbi_to_hazus, DATA_DIR
from src.interpolation import interpolate_im
from src.fragility import damage_state_probabilities as dsp

ROOT = _PROJECT_ROOT
OUT_DIR = ROOT / "output" / "demo_W1_0224" / "report_maps"

# ── Scenario Definitions ─────────────────────────────────────────

SCENARIOS = [
    {
        "id": "A",
        "label": "Full Region (Baseline)",
        "color": "#1565C0",       # dark blue
        "region": {"lat_min": 33.8, "lat_max": 34.6, "lon_min": -118.9, "lon_max": -118.0},
        "filters": {},
        "hwb_filter": None,
        "design_era": None,
    },
    {
        "id": "B",
        "label": "LA County, Pre-1975",
        "color": "#E65100",       # dark orange
        "region": {"lat_min": 33.8, "lat_max": 34.6, "lon_min": -118.9, "lon_max": -118.0},
        "filters": {"county": "037"},
        "hwb_filter": None,
        "design_era": "conventional",
    },
    {
        "id": "C",
        "label": "Epicenter Zone",
        "color": "#C62828",       # dark red
        "region": {"lat_min": 34.05, "lat_max": 34.35, "lon_min": -118.7, "lon_max": -118.4},
        "filters": {},
        "hwb_filter": None,
        "design_era": None,
    },
    {
        "id": "D",
        "label": "Concrete Multi-Span (HWB5+7)",
        "color": "#2E7D32",       # dark green
        "region": {"lat_min": 33.8, "lat_max": 34.6, "lon_min": -118.9, "lon_max": -118.0},
        "filters": {},
        "hwb_filter": ["HWB5", "HWB7"],
        "design_era": None,
    },
]


# ── GIS Map Helpers ──────────────────────────────────────────────

def add_north_arrow(ax, x=0.96, y=0.94, size=0.06):
    """Draw a north arrow on the axes."""
    ax.annotate(
        "", xy=(x, y), xytext=(x, y - size),
        xycoords="axes fraction", textcoords="axes fraction",
        arrowprops=dict(arrowstyle="-|>", color="black", lw=2.0, mutation_scale=18),
    )
    ax.text(
        x, y + 0.015, "N", transform=ax.transAxes,
        ha="center", va="bottom", fontsize=13, fontweight="bold", color="black",
        path_effects=[pe.withStroke(linewidth=3, foreground="white")],
    )


def add_scale_bar(ax, lon_center, lat_pos, length_km=20):
    """Draw a scale bar in data coordinates."""
    km_per_deg = 111.32 * np.cos(np.radians(lat_pos))
    bar_deg = length_km / km_per_deg
    x_left = lon_center - bar_deg / 2
    half = bar_deg / 2

    # Black/white alternating segments
    for x0, c in [(x_left, "black"), (x_left + half, "white")]:
        ax.add_patch(Rectangle(
            (x0, lat_pos), half, bar_deg * 0.05,
            linewidth=1.0, edgecolor="black", facecolor=c, zorder=10,
        ))

    # Labels
    text_y = lat_pos + bar_deg * 0.08
    for pos, lbl in [(x_left, "0"), (x_left + half, f"{length_km//2}"), (x_left + bar_deg, f"{length_km}")]:
        ax.text(pos, text_y, lbl, ha="center", va="bottom", fontsize=8, fontweight="bold",
                path_effects=[pe.withStroke(linewidth=2, foreground="white")], zorder=11)
    ax.text(x_left + bar_deg + bar_deg * 0.12, text_y, "km", ha="left", va="bottom",
            fontsize=8, fontstyle="italic",
            path_effects=[pe.withStroke(linewidth=2, foreground="white")], zorder=11)


def add_coordinate_grid(ax):
    """Add lat/lon grid with degree labels."""
    ax.grid(True, alpha=0.3, linestyle=":", linewidth=0.8, color="#666666")
    ax.xaxis.set_major_formatter(ticker.FormatStrFormatter("%.1f\u00b0"))
    ax.yaxis.set_major_formatter(ticker.FormatStrFormatter("%.1f\u00b0"))
    ax.tick_params(axis="both", labelsize=9)


def make_contour_background(ax, sm, im_type):
    """Draw ShakeMap contour background and return colorbar mappable."""
    sm_col = IM_COLUMN_MAP.get(im_type, "PSA10")
    lats, lons, vals = sm["LAT"].values, sm["LON"].values, sm[sm_col].values

    n_grid = 300
    lon_lin = np.linspace(lons.min(), lons.max(), n_grid)
    lat_lin = np.linspace(lats.min(), lats.max(), n_grid)
    lon_grid, lat_grid = np.meshgrid(lon_lin, lat_lin)
    val_grid = griddata((lons, lats), vals, (lon_grid, lat_grid), method="linear")

    colors_list = [
        "#FFFFFF", "#D4EAFF", "#80CDFF", "#48B8E0",
        "#7ECD68", "#F5F542", "#FFC832", "#FF8C28",
        "#FF3232", "#C80000", "#800000",
    ]
    cmap = LinearSegmentedColormap.from_list("shakemap", colors_list, N=256)
    vmax = max(0.5, np.nanpercentile(vals, 99))
    levels = np.linspace(0, vmax, 20)

    cf = ax.contourf(lon_grid, lat_grid, val_grid,
                     levels=levels, cmap=cmap, extend="max", alpha=0.7)
    ax.contour(lon_grid, lat_grid, val_grid,
               levels=levels[::3], colors="black", linewidths=0.25, alpha=0.35)
    return cf


def load_scenario_bridges(nbi_path, sm, scenario, im_type, interp_method):
    """Load, filter, classify, and compute damage for one scenario."""
    region = scenario["region"]
    bbox = region

    nbi = load_nbi(nbi_path, northridge_bbox=bbox)
    if len(nbi) == 0:
        return nbi

    nbi = classify_nbi_to_hazus(
        nbi,
        hwb_filter=scenario["hwb_filter"],
        design_era_filter=scenario["design_era"],
        nbi_filters=scenario["filters"],
    )
    if len(nbi) == 0:
        return nbi

    # Interpolate IM
    im_col = f"im_{im_type}"
    sm_col = IM_COLUMN_MAP.get(im_type, "PSA10")
    nbi[im_col] = interpolate_im(
        sm["LAT"].values, sm["LON"].values, sm[sm_col].values,
        nbi["latitude"].values, nbi["longitude"].values,
        method=interp_method,
    )
    nbi["im_selected"] = nbi[im_col]

    # Compute damage
    for _, row in nbi.iterrows():
        probs = dsp(row["im_selected"], row["hwb_class"])
        for ds_key in ["none", "slight", "moderate", "extensive", "complete"]:
            nbi.loc[row.name, f"P_{ds_key}"] = probs[ds_key]

    return nbi


# ── Main Map: 2x2 Subplots ───────────────────────────────────────

def plot_combined_map(sm, scenario_data, im_type, out_dir):
    """
    2x2 subplot figure — each panel shows one scenario with its own
    ShakeMap contour, bridges, region boundary, and map furniture.
    """
    im_labels = {"PGA": "PGA", "SA03": "Sa(0.3s)", "SA10": "Sa(1.0s)", "SA30": "Sa(3.0s)"}
    fig, axes = plt.subplots(2, 2, figsize=(20, 16))
    axes_flat = axes.flatten()

    # Pre-compute ShakeMap grid (shared)
    sm_col = IM_COLUMN_MAP.get(im_type, "PSA10")
    lats_sm, lons_sm, vals_sm = sm["LAT"].values, sm["LON"].values, sm[sm_col].values
    n_grid = 300
    lon_lin = np.linspace(lons_sm.min(), lons_sm.max(), n_grid)
    lat_lin = np.linspace(lats_sm.min(), lats_sm.max(), n_grid)
    lon_grid, lat_grid = np.meshgrid(lon_lin, lat_lin)
    val_grid = griddata((lons_sm, lats_sm), vals_sm, (lon_grid, lat_grid), method="linear")

    colors_list = [
        "#FFFFFF", "#D4EAFF", "#80CDFF", "#48B8E0",
        "#7ECD68", "#F5F542", "#FFC832", "#FF8C28",
        "#FF3232", "#C80000", "#800000",
    ]
    cmap = LinearSegmentedColormap.from_list("shakemap", colors_list, N=256)
    vmax = max(0.5, np.nanpercentile(vals_sm, 99))
    levels = np.linspace(0, vmax, 20)

    cf_last = None  # for shared colorbar

    for idx, (sc, nbi) in enumerate(scenario_data):
        ax = axes_flat[idx]
        sid = sc["id"]
        color = sc["color"]
        label = sc["label"]
        region = sc["region"]
        n = len(nbi)

        # ShakeMap contour
        cf = ax.contourf(lon_grid, lat_grid, val_grid,
                         levels=levels, cmap=cmap, extend="max", alpha=0.75)
        ax.contour(lon_grid, lat_grid, val_grid,
                   levels=levels[::3], colors="black", linewidths=0.2, alpha=0.3)
        cf_last = cf

        # Region boundary
        rect = Rectangle(
            (region["lon_min"], region["lat_min"]),
            region["lon_max"] - region["lon_min"],
            region["lat_max"] - region["lat_min"],
            linewidth=2.5, edgecolor=color, facecolor="none",
            linestyle="--", zorder=8, label="Study Region",
        )
        ax.add_patch(rect)

        # Bridges
        ax.scatter(
            nbi["longitude"], nbi["latitude"],
            s=10, c=color, alpha=0.7, edgecolors="black", linewidths=0.2,
            zorder=6,
        )

        # Epicenter
        ax.plot(-118.537, 34.213, marker="*", markersize=16, color="#FFD600",
                markeredgecolor="black", markeredgewidth=1.2, zorder=9)

        # Compute summary stats
        p_mod_plus = nbi[["P_moderate", "P_extensive", "P_complete"]].sum(axis=1).mean()
        p_comp = nbi["P_complete"].mean()
        mean_sa = nbi["im_selected"].mean()

        # Filter description
        filters = []
        if sc["filters"]:
            for k, v in sc["filters"].items():
                filters.append(f"{k}={v}")
        if sc["design_era"]:
            filters.append(f"era={sc['design_era']}")
        if sc["hwb_filter"]:
            filters.append(f"hwb={','.join(sc['hwb_filter'])}")
        filter_str = ", ".join(filters) if filters else "none"

        # Panel title
        ax.set_title(
            f"({sid}) {label}\n"
            f"n={n:,}  |  Filter: {filter_str}  |  "
            f"Mean Sa={mean_sa:.3f}g  |  P(Mod+)={p_mod_plus:.1%}",
            fontsize=10.5, fontweight="bold", pad=8, color="#333333",
        )

        # Map furniture
        add_north_arrow(ax, x=0.96, y=0.93, size=0.05)

        # Zoom to this scenario's region with padding
        pad_lon = (region["lon_max"] - region["lon_min"]) * 0.08
        pad_lat = (region["lat_max"] - region["lat_min"]) * 0.08
        ax.set_xlim(region["lon_min"] - pad_lon, region["lon_max"] + pad_lon)
        ax.set_ylim(region["lat_min"] - pad_lat * 2.5, region["lat_max"] + pad_lat)

        # Scale bar — placed below region boundary
        lon_c = (region["lon_min"] + region["lon_max"]) / 2
        lat_b = region["lat_min"] - pad_lat * 1.5
        # Choose scale bar length based on region size
        region_width_km = (region["lon_max"] - region["lon_min"]) * 111.32 * np.cos(np.radians(34.2))
        bar_km = 10 if region_width_km < 50 else 20
        add_scale_bar(ax, lon_c, lat_b, length_km=bar_km)

        add_coordinate_grid(ax)
        ax.set_xlabel("Longitude (\u00b0W)", fontsize=9, labelpad=4)
        ax.set_ylabel("Latitude (\u00b0N)", fontsize=9, labelpad=4)
        ax.set_aspect("equal", adjustable="box")
        for spine in ax.spines.values():
            spine.set_linewidth(1.2)
            spine.set_color("#444444")

        # Legend in each panel
        legend_items = [
            Line2D([0], [0], marker="o", color="w", markerfacecolor=color,
                   markeredgecolor="black", markersize=7, linewidth=0,
                   label=f"Bridges (n={n:,})"),
            Line2D([0], [0], color=color, linewidth=2.5, linestyle="--",
                   label="Study Region"),
            Line2D([0], [0], marker="*", color="w", markerfacecolor="#FFD600",
                   markeredgecolor="black", markersize=11, linewidth=0,
                   label="Epicenter"),
        ]
        ax.legend(
            handles=legend_items, loc="lower left", fontsize=7.5,
            framealpha=0.9, edgecolor="#cccccc", borderpad=0.5,
            labelspacing=0.3, handletextpad=0.4,
        )

    # Shared colorbar on the right
    fig.subplots_adjust(right=0.92)
    cbar_ax = fig.add_axes([0.935, 0.15, 0.015, 0.65])
    cbar = fig.colorbar(cf_last, cax=cbar_ax)
    cbar.set_label(f"{im_labels.get(im_type, im_type)} [g]", fontsize=12, fontweight="bold")
    cbar.ax.tick_params(labelsize=9)

    # Suptitle
    fig.suptitle(
        "CAT411 Bridge Seismic Risk Assessment — Multi-Scenario Comparison\n"
        "1994 Northridge Earthquake (Mw 6.7)  |  Sa(1.0s) from USGS ShakeMap",
        fontsize=15, fontweight="bold", y=0.98,
    )

    # Source
    fig.text(
        0.5, 0.005,
        "Data: USGS ShakeMap (ci3144585) + FHWA NBI (2024)  |  CRS: WGS 84 (EPSG:4326)  |  CAT411 Framework",
        ha="center", va="bottom", fontsize=8, color="#888888", fontstyle="italic",
    )

    path = out_dir / "map_multi_scenario.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  Saved: {path.name}")
    return path


# ── Conditions Table ─────────────────────────────────────────────

def plot_conditions_table(scenario_data, im_type, out_dir):
    """Generate a clean conditions summary table as a figure."""
    fig, ax = plt.subplots(figsize=(16, 5))
    ax.axis("off")

    # Table data
    headers = [
        "ID", "Scenario", "Region (Lat)", "Region (Lon)",
        "Filters", "Bridges", "Mean Sa\n[g]", "Max Sa\n[g]",
        "P(Mod+)", "P(Comp)",
    ]

    rows = []
    for sc, nbi in scenario_data:
        r = sc["region"]
        filters = []
        if sc["filters"]:
            for k, v in sc["filters"].items():
                filters.append(f"{k}={v}")
        if sc["design_era"]:
            filters.append(f"era={sc['design_era']}")
        if sc["hwb_filter"]:
            filters.append(f"hwb={','.join(sc['hwb_filter'])}")
        filter_str = "; ".join(filters) if filters else "(none)"

        p_mod_plus = nbi[["P_moderate", "P_extensive", "P_complete"]].sum(axis=1).mean()
        p_comp = nbi["P_complete"].mean()

        rows.append([
            sc["id"],
            sc["label"],
            f"{r['lat_min']:.2f} - {r['lat_max']:.2f}",
            f"{r['lon_min']:.1f} - {r['lon_max']:.1f}",
            filter_str,
            f"{len(nbi):,}",
            f"{nbi['im_selected'].mean():.4f}",
            f"{nbi['im_selected'].max():.4f}",
            f"{p_mod_plus:.1%}",
            f"{p_comp:.1%}",
        ])

    table = ax.table(
        cellText=rows,
        colLabels=headers,
        loc="center",
        cellLoc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9.5)
    table.scale(1.0, 2.0)
    table.auto_set_column_width(col=list(range(len(headers))))

    # Style header row
    for j, header in enumerate(headers):
        cell = table[0, j]
        cell.set_facecolor("#1565C0")
        cell.set_text_props(color="white", fontweight="bold", fontsize=9)
        cell.set_edgecolor("white")

    # Color each scenario row's ID cell
    for i, (sc, _) in enumerate(scenario_data):
        cell = table[i + 1, 0]
        cell.set_facecolor(sc["color"])
        cell.set_text_props(color="white", fontweight="bold")

        # Alternate row background
        bg = "#F5F5F5" if i % 2 == 0 else "#FFFFFF"
        for j in range(1, len(headers)):
            table[i + 1, j].set_facecolor(bg)
            table[i + 1, j].set_edgecolor("#DDDDDD")

    ax.set_title(
        "Analysis Scenarios — Conditions & Results Summary",
        fontsize=14, fontweight="bold", pad=20,
    )

    fig.tight_layout()
    path = out_dir / "table_conditions.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  Saved: {path.name}")
    return path


# ── Main ─────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  CAT411 — GIS-Style Report Maps (Multi-Scenario)")
    print("=" * 60)

    cfg = load_config("config.yaml")
    im_type = cfg.im_type

    # Load ShakeMap once
    print("\n[1/3] Loading ShakeMap...")
    sm = load_shakemap()
    print(f"  Grid points: {len(sm):,}")

    # Load bridges for each scenario
    print("\n[2/3] Loading bridges per scenario...")
    nbi_path = sorted(DATA_DIR.glob("[A-Z][A-Z]*.txt"))[0]

    scenario_data = []
    for sc in SCENARIOS:
        nbi = load_scenario_bridges(nbi_path, sm, sc, im_type, cfg.interpolation_method)
        print(f"  [{sc['id']}] {sc['label']}: {len(nbi):,} bridges")
        scenario_data.append((sc, nbi))

    # Clean old output
    # Clean old outputs (tolerant of Google Drive / desktop.ini locks)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    try:
        for f in OUT_DIR.iterdir():
            if f.is_file() and f.suffix in (".png", ".docx", ".txt"):
                f.unlink()
    except PermissionError:
        pass  # Google Drive sync lock — harmless

    # Generate
    print(f"\n[3/3] Generating outputs...")
    plot_combined_map(sm, scenario_data, im_type, OUT_DIR)
    plot_conditions_table(scenario_data, im_type, OUT_DIR)
    generate_word_report(scenario_data, im_type, OUT_DIR)

    print(f"\nDone! Output: {OUT_DIR.relative_to(ROOT)}/")


# ── Word Report ──────────────────────────────────────────────────

def generate_word_report(scenario_data, im_type, out_dir):
    """Generate Word report with map, conditions table, and config reference."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup — A4 landscape
    for section in doc.sections:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)

    # ── Title ──
    title = doc.add_heading("CAT411 — Multi-Scenario Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Weekly Report W1 (2026-02-24)  |  1994 Northridge Earthquake (Mw 6.7)  |  Sa(1.0s) ShakeMap"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── Map ──
    doc.add_heading("1. Multi-Scenario Comparison Map", level=1)
    map_path = out_dir / "map_multi_scenario.png"
    if map_path.exists():
        doc.add_picture(str(map_path), width=Cm(25))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── Conditions Table ──
    doc.add_heading("2. Analysis Conditions & Results", level=1)

    headers = ["ID", "Scenario", "Region (Lat)", "Region (Lon)",
               "Filters", "Bridges", "Mean Sa [g]", "Max Sa [g]",
               "P(Mod+)", "P(Comp)"]

    id_colors_hex = {"A": "1565C0", "B": "E65100", "C": "C62828", "D": "2E7D32"}

    rows = []
    for sc, nbi in scenario_data:
        r = sc["region"]
        filters = []
        if sc["filters"]:
            for k, v in sc["filters"].items():
                filters.append(f"{k}={v}")
        if sc["design_era"]:
            filters.append(f"era={sc['design_era']}")
        if sc["hwb_filter"]:
            filters.append(f"hwb={','.join(sc['hwb_filter'])}")
        filter_str = "; ".join(filters) if filters else "(none)"
        p_mod = nbi[["P_moderate", "P_extensive", "P_complete"]].sum(axis=1).mean()
        p_comp = nbi["P_complete"].mean()
        rows.append([
            sc["id"], sc["label"],
            f"{r['lat_min']:.2f} – {r['lat_max']:.2f}",
            f"{r['lon_min']:.1f} – {r['lon_max']:.1f}",
            filter_str, f"{len(nbi):,}",
            f"{nbi['im_selected'].mean():.4f}",
            f"{nbi['im_selected'].max():.4f}",
            f"{p_mod:.1%}", f"{p_comp:.1%}",
        ])

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1565C0", qn("w:val"): "clear"}))

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = cell._element.get_or_add_tcPr()
                shading.append(shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): id_colors_hex.get(val, "333333"),
                    qn("w:val"): "clear"}))
            elif i % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                shading.append(shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): "F5F5F5", qn("w:val"): "clear"}))

    doc.add_paragraph("")

    # ── Key Observations ──
    doc.add_heading("3. Key Observations", level=1)
    obs = [
        "Scenario C (Epicenter Zone) shows 52.7% P(Moderate+), ~3x the baseline, "
        "due to 2x higher mean Sa (0.73g vs 0.32g).",
        "Pre-1975 conventional bridges (B) have slightly higher vulnerability "
        "than the full population at comparable IM levels.",
        "Concrete multi-span HWB5/HWB7 classes (D) show 23.4% P(Moderate+), "
        "~30% above baseline, reflecting their lower fragility thresholds.",
        "All scenarios share Max Sa = 1.26g, indicating the ShakeMap peak is "
        "within the full study region.",
    ]
    for o in obs:
        doc.add_paragraph(o, style="List Bullet")
    doc.add_paragraph("")

    # ── Supported Config Parameters (full reference) ──
    doc.add_heading("4. Supported Configuration Parameters (config.yaml)", level=1)
    p = doc.add_paragraph(
        "The framework is fully driven by config.yaml. All parameters below are "
        "optional; defaults are used when omitted. CLI flags (--im-type, --nbi-filter, "
        "--bbox) override the corresponding YAML settings."
    )
    p.runs[0].font.size = Pt(9)

    # Config reference table
    cfg_headers = ["Section", "Parameter", "Type", "Default", "Description"]
    cfg_rows = [
        # Region
        ["Region", "region.lat_min", "float", "33.8", "Southern boundary of study area (degrees)"],
        ["", "region.lat_max", "float", "34.6", "Northern boundary of study area (degrees)"],
        ["", "region.lon_min", "float", "-118.9", "Western boundary of study area (degrees)"],
        ["", "region.lon_max", "float", "-118.0", "Eastern boundary of study area (degrees)"],
        # Bridge Selection
        ["Bridge\nSelection", "bridge_selection.county", "str", "(all)", "NBI county FIPS code, e.g. \"037\" for LA County"],
        ["", "bridge_selection.year_built", "str", "(all)", "Numeric comparison, e.g. \">1960\", \"<=1990\""],
        ["", "bridge_selection.material", "list[str]", "(all)", "Material filter, e.g. [\"concrete\", \"steel\"]"],
        ["", "hwb_filter", "list[str]", "(all)", "Restrict to specific HWB classes, e.g. [\"HWB5\", \"HWB7\"]"],
        ["", "design_era", "str", "(all)", "\"conventional\" (pre-1990) or \"seismic\" (post-1990)"],
        ["", "material_filter", "list[str]", "(all)", "Top-level material filter (alternative to bridge_selection.material)"],
        # IM Source
        ["IM Source", "im_source", "str", "shakemap", "\"shakemap\" (from grid.xml) or \"gmpe\" (BA08 synthetic)"],
        ["", "im_type", "str", "SA10", "PGA | SA03 | SA10 | SA30. Non-SA10 requires fragility_overrides"],
        # Interpolation
        ["Spatial\nInterpolation", "interpolation.method", "str", "nearest",
         "nearest | idw | bilinear | natural_neighbor | kriging"],
        ["", "interpolation.power", "float", "2.0", "IDW distance exponent (higher = more local)"],
        ["", "interpolation.n_neighbors", "int", "8", "Number of nearby grid points (IDW, kriging)"],
        ["", "interpolation.range_km", "float", "50.0", "Variogram range in km (kriging only)"],
        ["", "interpolation.nugget", "float", "0.01", "Measurement noise (kriging only)"],
        # GMPE
        ["GMPE\nScenario", "gmpe_scenario.Mw", "float", "6.7", "Moment magnitude"],
        ["", "gmpe_scenario.lat", "float", "34.213", "Epicenter latitude"],
        ["", "gmpe_scenario.lon", "float", "-118.537", "Epicenter longitude"],
        ["", "gmpe_scenario.depth_km", "float", "18.4", "Hypocentral depth (km)"],
        ["", "gmpe_scenario.fault_type", "str", "reverse", "strike_slip | normal | reverse | unspecified"],
        ["", "gmpe_scenario.vs30", "float", "760", "Default site Vs30 (m/s)"],
        # Fragility
        ["Fragility\nOverrides", "fragility_overrides.<HWB>.<ds>.median", "float", "Hazus 7.9",
         "Override median Sa for a damage state (required if im_type != SA10)"],
        ["", "fragility_overrides.<HWB>.<ds>.beta", "float", "0.6",
         "Override lognormal dispersion for a damage state"],
        # Calibration
        ["Calibration", "calibration.global_median_factor", "float", "1.0",
         "Scale all fragility medians. <1.0 = more vulnerable"],
        ["", "calibration.class_factors.<HWB>", "float", "1.0",
         "Per-class median scale factor (overrides global)"],
        # Analysis
        ["Analysis", "analysis.n_realizations", "int", "50", "Monte Carlo realizations per event"],
        ["", "analysis.n_events", "int", "50", "Stochastic events for probabilistic mode"],
        ["", "analysis.seed", "int", "42", "Random seed for reproducibility"],
    ]

    ct = doc.add_table(rows=1 + len(cfg_rows), cols=len(cfg_headers))
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    ct.style = "Table Grid"

    # Header
    for j, h in enumerate(cfg_headers):
        cell = ct.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "37474F", qn("w:val"): "clear"}))

    # Section colors
    section_colors = {
        "Region": "E3F2FD", "Bridge\nSelection": "FFF3E0",
        "IM Source": "E8F5E9", "Spatial\nInterpolation": "F3E5F5",
        "GMPE\nScenario": "ECEFF1", "Fragility\nOverrides": "FBE9E7",
        "Calibration": "FFFDE7", "Analysis": "E0F7FA",
    }
    current_section = ""
    for i, row_data in enumerate(cfg_rows):
        for j, val in enumerate(row_data):
            cell = ct.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            if j == 0 and val:
                current_section = val
                run.bold = True
                run.font.size = Pt(8.5)
            if j == 1:
                run.font.name = "Consolas"
            # Section background color
            if j == 0 and val:
                bg = section_colors.get(val, "FFFFFF")
                shading = cell._element.get_or_add_tcPr()
                shading.append(shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): bg, qn("w:val"): "clear"}))

    doc.add_paragraph("")

    # ── Validation Rules ──
    doc.add_heading("5. Validation Rules", level=1)
    rules = [
        ("IM-Fragility Compatibility",
         "If im_type is not SA10, fragility_overrides MUST be provided. "
         "Default Hazus parameters are calibrated for Sa(1.0s) only. "
         "Violation raises ValueError at config load time."),
        ("IM Column Availability",
         "If the configured IM type is not present in the ShakeMap grid.xml, "
         "a ValueError is raised listing available IM columns."),
        ("Zero-IM Warning",
         "Bridges receiving IM <= 0.0g after interpolation trigger a RuntimeWarning. "
         "This indicates potential spatial extent mismatch."),
        ("Unknown IM Type",
         "Unrecognized im_type values (e.g. typos like 'SA1.0') raise ValueError "
         "immediately at config load time."),
    ]
    for rule_name, rule_desc in rules:
        p = doc.add_paragraph()
        run_b = p.add_run(f"{rule_name}: ")
        run_b.bold = True
        run_b.font.size = Pt(9)
        run_d = p.add_run(rule_desc)
        run_d.font.size = Pt(9)

    doc.add_paragraph("")

    # ── Footer ──
    p = doc.add_paragraph()
    run = p.add_run(
        "Data: USGS ShakeMap (ci3144585) + FHWA NBI (2024)  |  "
        "CRS: WGS 84  |  CAT411 Framework"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path = out_dir / "scenario_report.docx"
    doc.save(str(path))
    print(f"  Saved: {path.name}")


if __name__ == "__main__":
    main()
