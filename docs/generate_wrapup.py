"""Generate the CAT411 Wrap-Up document (Word) with architecture flowchart."""

import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output" / "wrapup"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ── 1. Generate Architecture Flowchart ─────────────────────────────────────

def draw_flowchart(save_path: Path):
    """Draw the system architecture flowchart."""
    fig, ax = plt.subplots(figsize=(10, 7.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    # Color palette
    c_user = "#E3F2FD"      # light blue
    c_agent = "#FFF3E0"     # light orange
    c_llm = "#F3E5F5"       # light purple
    c_tool = "#E8F5E9"      # light green
    c_cat = "#FFEBEE"       # light red
    c_output = "#FFF9C4"    # light yellow
    border = "#37474F"

    box_style = dict(boxstyle="round,pad=0.4", linewidth=1.8, edgecolor=border)
    title_font = dict(fontsize=11, fontweight="bold", fontfamily="sans-serif")
    body_font = dict(fontsize=8.5, fontfamily="sans-serif", color="#424242")
    arrow_kw = dict(
        arrowstyle="->,head_width=0.3,head_length=0.2",
        color="#546E7A", linewidth=2,
    )
    label_font = dict(fontsize=7.5, fontstyle="italic", color="#546E7A",
                      fontfamily="sans-serif")

    # ── Boxes ──

    # 1) User
    ax.text(5, 9.3, "End User", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_user), **title_font)
    ax.text(5, 8.75, 'Natural language request\n"Estimate damage for M6.7 near LA\nusing real NBI bridges"',
            ha="center", va="center", **body_font)

    # 2) Agent Core
    ax.text(5, 7.1, "Agent Core  (Python CLI)", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_agent), **title_font)
    ax.text(5, 6.55, "Intent parsing · Prompt construction\nConversation history · Code assembly",
            ha="center", va="center", **body_font)

    # 3) LLM API
    ax.text(5, 5.0, "LLM API  (OpenAI / Anthropic)", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_llm), **title_font)
    ax.text(5, 4.45, "Function calling · Code generation\nParameter extraction from NL query",
            ha="center", va="center", **body_font)

    # 4) Code Generator (left side)
    ax.text(2.3, 3.1, "Code Generator", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_tool), **title_font)
    ax.text(2.3, 2.55, "Assemble executable\nPython script from\nLLM tool selections",
            ha="center", va="center", **body_font)

    # 5) CAT411 Tool Layer (right side)
    ax.text(7.7, 3.1, "CAT411 Tool Layer", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_tool), **title_font)
    ax.text(7.7, 2.4, "8 tools wrapping framework APIs:\nquery_bridges · run_scenario\ncompute_loss · plot_results\nget_fragility · summarize_portfolio",
            ha="center", va="center", **body_font)

    # 6) CAT411 Framework
    ax.text(5, 1.1, "CAT411 Framework  (Existing)", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_cat), **title_font)
    ax.text(5, 0.5, "Hazus 6.1 fragility · BA08/BSSA21 GMPE · 25,000+ CA bridges\nLoss computation · 3-level validation · Visualization",
            ha="center", va="center", **body_font)

    # 7) Output (far right)
    ax.text(9.3, 1.1, "Output", ha="center", va="center",
            bbox=dict(**box_style, facecolor=c_output), **title_font)
    ax.text(9.3, 0.55, "Executable\n.py script\n+ results",
            ha="center", va="center", **body_font)

    # ── Arrows ──

    # User → Agent
    ax.annotate("", xy=(5, 7.55), xytext=(5, 8.35), arrowprops=arrow_kw)
    ax.text(5.6, 7.95, "NL query", **label_font)

    # Agent → LLM
    ax.annotate("", xy=(5, 5.45), xytext=(5, 6.15), arrowprops=arrow_kw)
    ax.text(5.6, 5.8, "prompt + tool schemas", **label_font)

    # LLM → Code Generator
    ax.annotate("", xy=(2.3, 3.55), xytext=(4.0, 4.1),
                arrowprops=dict(**arrow_kw, connectionstyle="arc3,rad=0.2"))
    ax.text(2.2, 3.95, "code template\n+ parameters", **label_font)

    # LLM → Tool Layer
    ax.annotate("", xy=(7.7, 3.55), xytext=(6.0, 4.1),
                arrowprops=dict(**arrow_kw, connectionstyle="arc3,rad=-0.2"))
    ax.text(7.0, 3.95, "tool calls", **label_font)

    # Code Generator → CAT411
    ax.annotate("", xy=(3.8, 1.5), xytext=(2.3, 2.1),
                arrowprops=dict(**arrow_kw, connectionstyle="arc3,rad=-0.15"))

    # Tool Layer → CAT411
    ax.annotate("", xy=(6.2, 1.5), xytext=(7.7, 2.0),
                arrowprops=dict(**arrow_kw, connectionstyle="arc3,rad=0.15"))

    # CAT411 → Output
    ax.annotate("", xy=(8.5, 1.1), xytext=(6.8, 1.1), arrowprops=arrow_kw)

    # Output → User (dashed return)
    ax.annotate("", xy=(8.8, 9.0), xytext=(9.3, 1.55),
                arrowprops=dict(
                    arrowstyle="->,head_width=0.25,head_length=0.18",
                    color="#90A4AE", linewidth=1.5, linestyle="dashed",
                    connectionstyle="arc3,rad=-0.3",
                ))
    ax.text(9.6, 5.0, "script +\nresults", rotation=90, **label_font)

    # ── Title ──
    ax.text(5, 9.85, "System Architecture: LLM Code-Generation Agent for Seismic Bridge Risk Assessment",
            ha="center", va="center", fontsize=12, fontweight="bold",
            fontfamily="sans-serif", color="#1A237E")

    # ── One-time cost badge ──
    ax.text(5, 4.0, "LLM called once →  generates reusable script  → zero ongoing cost",
            ha="center", va="center", fontsize=8, fontweight="bold",
            color="#E65100",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#FFF3E0",
                      edgecolor="#E65100", linewidth=1.2, alpha=0.9))

    plt.tight_layout()
    fig.savefig(str(save_path), dpi=200, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    print(f"Flowchart saved to: {save_path}")


# ── 2. Generate Word Document ──────────────────────────────────────────────

def create_wrapup_doc(flowchart_path: Path, save_path: Path):
    """Create the 2-page wrap-up Word document."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── Title ──
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "LLM Code-Generation Agent for Seismic Bridge Risk Assessment"
    )
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("CAT411 Framework Wrap-Up Project")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

    # ── 1. Introduction & Motivation ──
    doc.add_heading("1. Introduction & Motivation", level=1)

    doc.add_paragraph(
        "Seismic risk assessment for transportation infrastructure requires "
        "specialized knowledge of ground motion prediction equations (GMPEs), "
        "structural fragility models, and loss estimation methodologies. "
        "The CAT411 framework implements the FEMA Hazus 6.1 earthquake model "
        "for bridge damage and loss estimation, supporting 25,000+ California "
        "bridges with dual hazard paths (USGS ShakeMap interpolation and "
        "BSSA21 NGA-West2 GMPE), 28 Hazus bridge fragility classes, and a "
        "three-level validation framework against 1994 Northridge earthquake "
        "observations."
    )

    doc.add_paragraph(
        "However, using this framework requires programming expertise in Python "
        "and domain knowledge of seismic engineering parameters\u2014creating a "
        "significant barrier for practitioners, emergency managers, and "
        "decision-makers who need rapid risk estimates. Meanwhile, Large "
        "Language Models (LLMs) have demonstrated strong capabilities in "
        "natural language understanding and code generation, particularly "
        "through function calling (tool use) interfaces."
    )

    doc.add_paragraph(
        "This project bridges the gap by building an LLM-powered "
        "code-generation agent that translates natural language queries into "
        "executable Python scripts calling the CAT411 framework. Unlike "
        "traditional chatbot approaches that require continuous API calls, "
        "our agent uses the LLM only once to understand the user\u2019s intent "
        "and generate a complete, reusable script\u2014minimizing ongoing costs "
        "while maximizing accessibility."
    )

    # ── 2. Proposed Method ──
    doc.add_heading("2. Proposed Method", level=1)

    doc.add_heading("2.1 System Architecture", level=2)
    doc.add_paragraph(
        "The system follows a five-stage pipeline (Figure 1): "
        "(1) the user describes their analysis need in natural language via "
        "a CLI interface; "
        "(2) the Agent Core constructs a structured prompt embedding the "
        "user\u2019s request alongside tool definitions (JSON schemas) for "
        "8 CAT411 wrapper functions; "
        "(3) the LLM API (OpenAI GPT-4o or Anthropic Claude) performs intent "
        "extraction and selects the appropriate tool sequence with parameters; "
        "(4) the Code Generator assembles an executable Python script by "
        "composing the selected tool calls with the extracted parameters; "
        "(5) the user executes the generated script locally, which calls "
        "the CAT411 framework directly\u2014no further LLM interaction required."
    )

    # Insert flowchart
    doc.add_picture(str(flowchart_path), width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph("Figure 1. System architecture of the LLM code-generation agent.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True

    doc.add_heading("2.2 Tool Layer Design", level=2)
    doc.add_paragraph(
        "Eight tools wrap the core CAT411 APIs, each defined with a JSON "
        "schema describing its parameters, types, and defaults:"
    )

    # Tool table
    table = doc.add_table(rows=9, cols=3)
    table.style = "Light Shading Accent 1"
    headers = ["Tool Name", "CAT411 Module", "Function"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    tools_data = [
        ("query_bridges", "data_loader, bridge_classes", "Search NBI inventory by region/class/material"),
        ("run_scenario", "engine, hazard, exposure", "Deterministic earthquake scenario analysis"),
        ("compute_bridge_loss", "loss, fragility", "Single-bridge expected loss computation"),
        ("get_fragility", "hazus_params, fragility", "Fragility parameters and damage probabilities"),
        ("plot_fragility_curves", "plotting", "Generate fragility curve visualizations"),
        ("plot_damage_distribution", "plotting", "Damage state bar charts at multiple Sa levels"),
        ("plot_class_comparison", "plotting", "Cross-class fragility comparison plots"),
        ("summarize_portfolio", "exposure, data_loader", "Aggregate portfolio statistics for a region"),
    ]
    for row_idx, (name, module, func) in enumerate(tools_data, start=1):
        table.rows[row_idx].cells[0].text = name
        table.rows[row_idx].cells[1].text = module
        table.rows[row_idx].cells[2].text = func
        for cell in table.rows[row_idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()  # spacing

    doc.add_heading("2.3 Code Generation Strategy", level=2)
    doc.add_paragraph(
        "Rather than executing tools in real time and returning natural "
        "language summaries (which requires an LLM call for every interaction), "
        "our agent generates a standalone Python script that: "
        "(a) imports the necessary CAT411 modules, "
        "(b) configures parameters extracted from the user\u2019s query, "
        "(c) executes the analysis pipeline, and "
        "(d) saves results (tables, plots, summaries) to the output directory. "
        "The generated script is self-contained and can be re-executed with "
        "modified parameters without any LLM involvement, reducing the "
        "per-analysis cost to zero after the initial generation."
    )

    # ── 3. Expected Results ──
    doc.add_heading("3. Expected Results", level=1)

    doc.add_paragraph(
        "We anticipate the following outcomes from this project:"
    )

    bullets = [
        "Functional CLI agent capable of translating natural language "
        "earthquake risk queries into executable CAT411 scripts, validated "
        "against 5+ standard analysis scenarios including the 1994 Northridge "
        "earthquake (M6.7).",

        "Demonstrated cost efficiency: a single LLM API call ($0.01\u2013$0.05) "
        "produces a reusable script, compared to $0.10\u2013$0.50 per session "
        "for continuous-conversation agents.",

        "Quantitative correctness verification: generated scripts produce "
        "results identical to direct API usage (loss within \u00b10.1% of "
        "reference values for the same scenario parameters).",

        "User accessibility improvement: non-programmer users can obtain "
        "bridge damage estimates, fragility analyses, and portfolio risk "
        "summaries through plain English descriptions, lowering the barrier "
        "to seismic risk assessment.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)

    # ── 4. Timeline ──
    doc.add_heading("4. Implementation Timeline", level=1)

    timeline = doc.add_table(rows=5, cols=3)
    timeline.style = "Light Shading Accent 1"
    t_headers = ["Week", "Task", "Deliverable"]
    for i, h in enumerate(t_headers):
        cell = timeline.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)

    t_data = [
        ("1", "Agent core + 4 basic tools", "CLI with query, fragility, loss, scenario tools"),
        ("2", "Code generation + 4 advanced tools", "Script generator + plot/validation/compare tools"),
        ("3", "Integration testing + optimization", "End-to-end validation, Notebook demo"),
        ("4", "Report writing + demo preparation", "Technical report + presentation materials"),
    ]
    for row_idx, (w, task, deliv) in enumerate(t_data, start=1):
        timeline.rows[row_idx].cells[0].text = w
        timeline.rows[row_idx].cells[1].text = task
        timeline.rows[row_idx].cells[2].text = deliv
        for cell in timeline.rows[row_idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # ── Save ──
    doc.save(str(save_path))
    print(f"Document saved to: {save_path}")


# ── Main ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    flowchart_path = OUTPUT_DIR / "architecture_flowchart.png"
    doc_path = OUTPUT_DIR / "CAT411_WrapUp_Proposal.docx"

    draw_flowchart(flowchart_path)
    create_wrapup_doc(flowchart_path, doc_path)
    print("\nDone!")
